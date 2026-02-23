from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from fuzzywuzzy import process, fuzz
from tkinter import ttk, filedialog, messagebox, StringVar, Listbox, MULTIPLE, END, Tk, Entry, Button
import glob
import json
import logging
import os
import pandas as pd
import re
import requests
import subprocess
import tkinter as tk
import threading
import pythoncom
import win32com.client as win32
import time

# Константы Word для печати
WD_PRINT_ALL_DOCUMENT = 0
WD_PRINT_SELECTION = 2
WD_PRINT_CURRENT_PAGE = 3
WD_PRINT_RANGE_OF_PAGES = 4
WD_PRINT_PRINT_TO_FILE = 5

# Отключаем логирование для win32com
logging.getLogger('win32com').setLevel(logging.WARNING)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename=os.path.join(os.getcwd(), 'log.txt'),
    filemode='w'
)

def extract_text_from_docx(file_path):
    """Извлекает текст из файла .docx."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return full_text

def extract_text_from_doc(file_path):
    """Извлекает текст из файла .doc."""
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(file_path)
    full_text = []
    for para in doc.Paragraphs:
        full_text.append(para.Range.Text.strip())
    doc.Close()
    word.Quit()
    return full_text

def extract_info(text_lines):
    info = {'Наименование': '', 'ИНН': '', 'Адрес': '', 'Электронная почта': '', 'Телефон': ''}
    
    # Убираем пустые строки
    text_lines = [line for line in text_lines if line.strip()]
    
    # Останавливаем обработку при обнаружении слова "ЗАПРОС" (в любом регистре)
    stop_index = None
    for i, line in enumerate(text_lines):
        if "ЗАПРОС" in line.upper():
            stop_index = i
            break
    
    # Если слово "ЗАПРОС" не найдено, ищем три пустые строки подряд
    if stop_index is None:
        empty_lines_count = 0
        for i, line in enumerate(text_lines):
            if not line.strip():
                empty_lines_count += 1
                if empty_lines_count == 3:
                    stop_index = i - 2
                    break
            else:
                empty_lines_count = 0
    
    # Если найдено условие остановки, обрезаем текст
    if stop_index is not None:
        text_lines = text_lines[:stop_index]
    
    combined_text = ' '.join(text_lines)

    # Поиск наименования
    inn_index = next((i for i, line in enumerate(text_lines) if 'ИНН' in line), None)
    if inn_index is not None:
        info['Наименование'] = ' '.join(line.strip() for line in text_lines[:inn_index])
    else:
        info['Наименование'] = text_lines[0].strip() if text_lines else ''

    # Удаляем обращения
    info['Наименование'] = re.sub(
        r'^(Руководителю|ИП|Индивидуальный предприниматель|Индивидуальному предпринимателю|Директору|Генеральному директору)\s*',
        '', info['Наименование'], flags=re.IGNORECASE
    ).strip()

    # Если наименование начинается с кавычек, добавляем "ООО"
    if info['Наименование'].startswith(('«', '"', "'")):
        info['Наименование'] = f"ООО {info['Наименование']}"

    # Заменяем "Общество с ограниченной ответственностью" на "ООО"
    info['Наименование'] = info['Наименование'].replace("Общество с ограниченной ответственностью", "ООО")

    # Поиск ИНН
    inn_match = re.search(r'ИНН\s*(\d{10,12})', combined_text)
    if inn_match:
        info['ИНН'] = inn_match.group(1).strip()

    # Поиск адреса
    address_match = re.search(r'(?:Адрес|Юридический адрес):\s*([\s\S]+?)(?=\sE-mail|\sТелефон|$)', combined_text)
    if address_match:
        info['Адрес'] = address_match.group(1).strip()

    # Поиск всех адресов электронной почты
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', combined_text)
    if emails:
        info['Электронная почта'] = ', '.join(emails)

    # Поиск всех номеров телефонов
    phones = re.findall(r'\+?\d[\d\s\-()]{6,}\d', combined_text)
    if phones:
        info['Телефон'] = ', '.join(phones)

    # Сохраняем всю информацию до последней строки с email
    info['Исходная информация'] = '\n'.join(text_lines)

    return info

def get_inn_by_name(organization_name):
    """Поиск ИНН через Яндекс."""
    try:
        search_query = f"ИНН {organization_name}"
        url = f"https://yandex.ru/search/?text={requests.utils.quote(search_query)}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')
        search_results = soup.find_all('li', class_='serp-item')
        for result in search_results:
            text = result.get_text()
            inn_match = re.search(r'\b\d{10,12}\b', text)
            if inn_match:
                return inn_match.group(0)

        logging.warning(f"ИНН для организации '{organization_name}' не найден в поисковой выдаче Яндекс.")
        return None
    except Exception as e:
        logging.error(f"Ошибка при поиске ИНН для организации '{organization_name}': {e}")
        return None

def create_output_file(progress_bar, status_label):
    """Создает файл output.xlsx на основе данных из документов."""
    current_dir = working_folder_var.get()
    if not current_dir:
        messagebox.showerror("Ошибка", "Выберите рабочую папку.")
        return

    doc_files = glob.glob(os.path.join(current_dir, '*.doc')) + glob.glob(os.path.join(current_dir, '*.docx'))
    if not doc_files:
        messagebox.showerror("Ошибка", "В папке нет файлов .doc или .docx.")
        return

    data = []
    total_files = len(doc_files)
    for idx, file_path in enumerate(doc_files):
        try:
            if file_path.endswith('.docx'):
                text_lines = extract_text_from_docx(file_path)
            else:
                text_lines = extract_text_from_doc(file_path)

            info = extract_info(text_lines)

            # Удаляем строку, если наименование начинается с "Запрос", "Добрый", "ЕИС", "Единая"
            if info['Наименование'].lower().startswith(('запрос', 'добрый', 'еис', 'единая')):
                continue

            info['Номер п/п'] = idx + 1

            # Если ИНН не найден или совпадает с указанным значением, ищем через Яндекс
            if compare_inn_var.get() and (not info['ИНН'] or info['ИНН'] == compare_inn_value_var.get()):
                inn = get_inn_by_name(info['Наименование'])
                if inn:
                    info['ИНН'] = inn

            data.append(info)
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {file_path}: {e}")

        # Обновляем прогресс
        progress = (idx + 1) / total_files * 100
        progress_bar['value'] = progress
        status_label.config(text=f"Обработано {idx + 1} из {total_files} файлов")
        root.update_idletasks()

    # Создаем DataFrame и сохраняем в Excel
    df = pd.DataFrame(data)
    columns_order = ['Номер п/п', 'Наименование', 'ИНН', 'Адрес', 'Электронная почта', 'Телефон', 'Исходная информация']
    df = df[columns_order]
    output_path = os.path.join(os.getcwd(), 'output.xlsx')
    df.to_excel(output_path, index=False)
    messagebox.showinfo("Успех", f"Файл output.xlsx успешно создан: {output_path}")

def process_files(progress_bar, status_label):
    """Обрабатывает файлы и создает документы на основе шаблона."""
    output_file = output_file_var.get()
    template_file = template_file_var.get()
    working_folder = working_folder_var.get()
    replace_from = replace_from_var.get()
    replace_to = replace_to_var.get()

    if not output_file or not template_file or not working_folder:
        messagebox.showerror("Ошибка", "Выберите рабочую папку, файл output.xlsx и шаблон.")
        return

    try:
        df = pd.read_excel(output_file)
        required_columns = ['Наименование', 'ИНН', 'Электронная почта']
        if not all(column in df.columns for column in required_columns):
            messagebox.showerror("Ошибка", f"В файле {output_file} отсутствуют необходимые колонки: {required_columns}")
            return

        output_folder = os.path.join(working_folder, "Итоговые_документы")
        os.makedirs(output_folder, exist_ok=True)

        total_rows = len(df)
        for index, row in df.iterrows():
            name = row['Наименование']
            inn = row['ИНН']
            email = row['Электронная почта']

            doc = Document(template_file)
            for para in doc.paragraphs:
                if replace_from and replace_to:
                    para.text = para.text.replace(replace_from, replace_to)
                if name in para.text:
                    para.text = para.text.replace(name, name)
                if inn and f"ИНН: {inn}" in para.text:
                    para.text = para.text.replace(f"ИНН: {inn}", f"ИНН: {inn}")
                if email and f"E-mail: {email}" in para.text:
                    para.text = para.text.replace(f"E-mail: {email}", f"E-mail: {email}")

            file_name = re.sub(r'[\\/:*?"<>|]', '_', name)[:50]
            file_path = os.path.join(output_folder, f"{file_name}.docx")
            doc.save(file_path)

            progress = (index + 1) / total_rows * 100
            progress_bar['value'] = progress
            status_label.config(text=f"Обработано {index + 1} из {total_rows} строк")
            root.update_idletasks()

        messagebox.showinfo("Успех", "Документы успешно созданы!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

def select_working_folder():
    """Выбор рабочей папки."""
    folder = filedialog.askdirectory()
    if folder:
        working_folder_var.set(folder)

def select_output_file():
    """Выбор файла output.xlsx."""
    file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file:
        output_file_var.set(file)

def select_template_file():
    """Выбор файла шаблона."""
    file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file:
        template_file_var.set(file)

def is_word_installed():
    try:
        word = win32.Dispatch("Word.Application")
        word.Quit()
        return True
    except:
        return False


def print_first_page_vbs(file_path, progress_callback=None):
    """Печатает первую страницу через VBS скрипт (отдельный процесс)"""
    try:
        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            raise FileNotFoundError(f"Файл не найден: {abs_path}")
        
        # Создаем уникальное имя для VBS файла
        timestamp = int(time.time() * 1000)
        safe_filename = re.sub(r'[^a-zA-Z0-9]', '_', os.path.basename(file_path))
        vbs_path = os.path.join(os.environ['TEMP'], f'print_{timestamp}_{safe_filename}.vbs')
        
        # VBS скрипт для печати первой страницы
        vbs_script = f'''
' VBS скрипт для печати первой страницы
On Error Resume Next

Dim Word, Doc
Set Word = CreateObject("Word.Application")
Word.Visible = False
Word.DisplayAlerts = False

' Открываем документ
Set Doc = Word.Documents.Open("{abs_path}")

If Err.Number <> 0 Then
    WScript.Echo "Ошибка открытия: " & Err.Description
    Word.Quit
    WScript.Quit 1
End If

' Печатаем первую страницу
' Range=4 (wdPrintRangeOfPages), Pages="1"
Doc.PrintOut False, , , , , , , , , , , , , , 4, "1"

' Ждем отправки на печать
WScript.Sleep 2000

Doc.Close False
Word.Quit

WScript.Quit 0
'''
        
        # Сохраняем VBS скрипт
        with open(vbs_path, 'w', encoding='cp1251') as f:
            f.write(vbs_script)
        
        # Запускаем VBS скрипт в отдельном процессе (не ждем завершения)
        subprocess.Popen(
            ['cscript', '//nologo', vbs_path],
            creationflags=subprocess.CREATE_NO_WINDOW,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        
        # Сразу сообщаем об успехе
        if progress_callback:
            progress_callback(True, os.path.basename(file_path))
        
        # Планируем удаление VBS файла через 30 секунд
        def delete_vbs():
            time.sleep(30)
            try:
                if os.path.exists(vbs_path):
                    os.remove(vbs_path)
            except:
                pass
        
        threading.Thread(target=delete_vbs, daemon=True).start()
        
        return True
        
    except Exception as e:
        error_msg = f"Ошибка печати {os.path.basename(file_path)}: {str(e)}"
        logging.error(error_msg)
        if progress_callback:
            progress_callback(False, error_msg)
        return False

def print_all_document_vbs(file_path, progress_callback=None):
    """Печатает весь документ через VBS скрипт (отдельный процесс)"""
    try:
        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            raise FileNotFoundError(f"Файл не найден: {abs_path}")
        
        # Создаем уникальное имя для VBS файла
        timestamp = int(time.time() * 1000)
        safe_filename = re.sub(r'[^a-zA-Z0-9]', '_', os.path.basename(file_path))
        vbs_path = os.path.join(os.environ['TEMP'], f'print_all_{timestamp}_{safe_filename}.vbs')
        
        # VBS скрипт для печати всего документа
        vbs_script = f'''
' VBS скрипт для печати всего документа
On Error Resume Next

Dim Word, Doc
Set Word = CreateObject("Word.Application")
Word.Visible = False
Word.DisplayAlerts = False

' Открываем документ
Set Doc = Word.Documents.Open("{abs_path}")

If Err.Number <> 0 Then
    WScript.Echo "Ошибка открытия: " & Err.Description
    Word.Quit
    WScript.Quit 1
End If

' Печатаем весь документ
' Range=0 (wdPrintAllDocument)
Doc.PrintOut False, , 0

' Ждем отправки на печать
WScript.Sleep 2000

Doc.Close False
Word.Quit

WScript.Quit 0
'''
        
        # Сохраняем VBS скрипт
        with open(vbs_path, 'w', encoding='cp1251') as f:
            f.write(vbs_script)
        
        # Запускаем VBS скрипт в отдельном процессе (не ждем завершения)
        subprocess.Popen(
            ['cscript', '//nologo', vbs_path],
            creationflags=subprocess.CREATE_NO_WINDOW,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        
        # Сразу сообщаем об успехе
        if progress_callback:
            progress_callback(True, os.path.basename(file_path))
        
        # Планируем удаление VBS файла через 30 секунд
        def delete_vbs():
            time.sleep(30)
            try:
                if os.path.exists(vbs_path):
                    os.remove(vbs_path)
            except:
                pass
        
        threading.Thread(target=delete_vbs, daemon=True).start()
        
        return True
        
    except Exception as e:
        error_msg = f"Ошибка печати {os.path.basename(file_path)}: {str(e)}"
        logging.error(error_msg)
        if progress_callback:
            progress_callback(False, error_msg)
        return False

def print_first_pages():
    """Печать ТОЛЬКО ПЕРВОЙ СТРАНИЦЫ каждого документа"""
    folder = print_folder_var.get()
    if not folder:
        messagebox.showwarning("Ошибка", "Выберите папку с документами")
        return
    
    try:
        files = []
        for ext in ['*.doc', '*.docx']:
            files.extend(glob.glob(os.path.join(folder, ext)))
        
        if not files:
            messagebox.showinfo("Информация", "В папке нет документов Word")
            return
        
        files.sort()
        
        if not messagebox.askyesno("Подтверждение", 
                                   f"Будет напечатана ТОЛЬКО ПЕРВАЯ СТРАНИЦА {len(files)} документов.\n"
                                   "Продолжить?"):
            return
        
        # Создаем окно прогресса
        progress_window = tk.Toplevel(root)
        progress_window.title("Печать первых страниц")
        progress_window.geometry("450x200")
        progress_window.resizable(False, False)
        
        # Центрируем
        progress_window.update_idletasks()
        x = (progress_window.winfo_screenwidth() - 450) // 2
        y = (progress_window.winfo_screenheight() - 200) // 2
        progress_window.geometry(f"+{x}+{y}")
        
        # Элементы окна
        label = tk.Label(progress_window, text="Подготовка к печати первых страниц...", font=('Arial', 10, 'bold'))
        label.pack(pady=10)
        
        progress = ttk.Progressbar(progress_window, length=350, maximum=len(files))
        progress.pack(pady=10)
        
        status = tk.Label(progress_window, text="", font=('Arial', 9))
        status.pack(pady=5)
        
        counter = tk.Label(progress_window, text=f"0 из {len(files)}", font=('Arial', 9))
        counter.pack(pady=5)
        
        def print_worker():
            completed = 0
            failed = 0
            
            for i, file_path in enumerate(files):
                try:
                    abs_path = os.path.abspath(file_path)
                    # Создаем VBS скрипт с правильными параметрами как в рабочем макросе
                    vbs_script = f'''
On Error Resume Next
Dim Word, Doc
Set Word = CreateObject("Word.Application")
Word.Visible = False
Word.DisplayAlerts = False
' Открываем документ с параметрами ReadOnly и AddToRecentFiles как в макросе
Set Doc = Word.Documents.Open("{abs_path}", , True, True)
If Err.Number = 0 Then
    ' Печатаем только первую страницу с правильными параметрами
    Doc.PrintOut False, , 4, , , , , , "1"
    WScript.Sleep 2000
    Doc.Close False
End If
Word.Quit
WScript.Quit 0
'''
                    # Сохраняем VBS файл
                    vbs_path = os.path.join(os.environ['TEMP'], f'print_first_{i}_{int(time.time())}.vbs')
                    with open(vbs_path, 'w', encoding='cp1251') as f:
                        f.write(vbs_script)
                    
                    # Запускаем VBS скрипт
                    subprocess.Popen(
                        ['cscript', '//nologo', vbs_path],
                        creationflags=subprocess.CREATE_NO_WINDOW,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL
                    )
                    
                    completed += 1
                    status.config(text=f"✓ {os.path.basename(file_path)}")
                    
                    # Планируем удаление VBS файла
                    def del_vbs():
                        time.sleep(10)
                        try:
                            if os.path.exists(vbs_path):
                                os.remove(vbs_path)
                        except:
                            pass
                    threading.Thread(target=del_vbs, daemon=True).start()
                    
                except Exception as e:
                    failed += 1
                    status.config(text=f"✗ {os.path.basename(file_path)}")
                    logging.error(f"Ошибка печати {file_path}: {str(e)}")
                
                # Обновляем прогресс
                progress['value'] = i + 1
                counter.config(text=f"{i + 1} из {len(files)} (успешно: {completed}, ошибок: {failed})")
                progress_window.update()
                
                # Пауза между файлами
                time.sleep(3)
            
            # Завершение
            label.config(text="Печать первых страниц завершена!")
            status.config(text=f"Готово: {completed} файлов, ошибок: {failed}")
            tk.Button(progress_window, text="Закрыть", 
                     command=progress_window.destroy,
                     width=10).pack(pady=10)
        
        # Запускаем в отдельном потоке
        thread = threading.Thread(target=print_worker, daemon=True)
        thread.start()
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")
        logging.error(f"Ошибка в print_first_pages: {str(e)}")

def print_all_documents():
    """Печать ВСЕХ СТРАНИЦ каждого документа"""
    folder = print_folder_var.get()
    if not folder:
        messagebox.showwarning("Ошибка", "Выберите папку с документами")
        return
    
    try:
        # Получаем все doc и docx файлы
        files = []
        for ext in ['*.doc', '*.docx']:
            files.extend(glob.glob(os.path.join(folder, ext)))
        
        if not files:
            messagebox.showinfo("Информация", "В папке нет документов Word")
            return
        
        files.sort()
        
        # Подтверждение
        if not messagebox.askyesno("Подтверждение", 
                                   f"Будут напечатаны ВСЕ СТРАНИЦЫ {len(files)} документов.\n"
                                   "Продолжить?"):
            return
        
        # Создаем окно прогресса
        progress_window = tk.Toplevel(root)
        progress_window.title("Печать всех страниц")
        progress_window.geometry("450x200")
        progress_window.resizable(False, False)
        
        # Центрируем
        progress_window.update_idletasks()
        x = (progress_window.winfo_screenwidth() - 450) // 2
        y = (progress_window.winfo_screenheight() - 200) // 2
        progress_window.geometry(f"+{x}+{y}")
        
        # Элементы окна
        label = tk.Label(progress_window, text="Подготовка к печати всех страниц...", font=('Arial', 10, 'bold'))
        label.pack(pady=10)
        
        progress = ttk.Progressbar(progress_window, length=350, maximum=len(files))
        progress.pack(pady=10)
        
        status = tk.Label(progress_window, text="", font=('Arial', 9))
        status.pack(pady=5)
        
        counter = tk.Label(progress_window, text=f"0 из {len(files)}", font=('Arial', 9))
        counter.pack(pady=5)
        
        def print_worker():
            completed = 0
            failed = 0
            
            for i, file_path in enumerate(files):
                try:
                    # Используем os.startfile для печати всего документа
                    os.startfile(file_path, "print")
                    
                    completed += 1
                    status.config(text=f"✓ {os.path.basename(file_path)}")
                    
                except Exception as e:
                    failed += 1
                    status.config(text=f"✗ {os.path.basename(file_path)}")
                    logging.error(f"Ошибка печати {file_path}: {str(e)}")
                
                # Обновляем прогресс
                progress['value'] = i + 1
                counter.config(text=f"{i + 1} из {len(files)} (успешно: {completed}, ошибок: {failed})")
                progress_window.update()
                
                # Пауза между файлами
                time.sleep(2)
            
            # Завершение
            label.config(text="Печать всех страниц завершена!")
            status.config(text=f"Готово: {completed} файлов, ошибок: {failed}")
            tk.Button(progress_window, text="Закрыть", 
                     command=progress_window.destroy,
                     width=10).pack(pady=10)
        
        # Запускаем в отдельном потоке
        thread = threading.Thread(target=print_worker, daemon=True)
        thread.start()
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")
        logging.error(f"Ошибка в print_all_documents: {str(e)}")

class PrintProgressWindow:
    """Окно прогресса печати"""
    def __init__(self, parent, total_files):
        self.top = tk.Toplevel(parent)
        self.top.title("Печать документов")
        self.top.geometry("450x200")
        self.top.resizable(False, False)
        
        # Центрируем окно
        self.top.update_idletasks()
        x = (self.top.winfo_screenwidth() - self.top.winfo_width()) // 2
        y = (self.top.winfo_screenheight() - self.top.winfo_height()) // 2
        self.top.geometry(f"+{x}+{y}")
        
        # Заголовок
        self.label = tk.Label(self.top, text="Подготовка к печати...", font=('Arial', 10, 'bold'))
        self.label.pack(pady=10)
        
        # Прогресс бар
        self.progress = ttk.Progressbar(self.top, length=400, maximum=total_files)
        self.progress.pack(pady=10, padx=20)
        
        # Статус
        self.status = tk.Label(self.top, text="", font=('Arial', 9))
        self.status.pack(pady=5)
        
        # Счетчик
        self.counter_label = tk.Label(self.top, text=f"0 из {total_files}", font=('Arial', 9))
        self.counter_label.pack(pady=5)
        
        # Кнопка закрытия (изначально неактивна)
        self.close_btn = tk.Button(
            self.top, 
            text="Закрыть", 
            command=self.close,
            state=tk.DISABLED,
            width=10
        )
        self.close_btn.pack(pady=10)
        
        self.completed = 0
        self.failed = 0
        self.total = total_files
        self.is_complete = False
    
    def update(self, success, message):
        """Обновляет прогресс"""
        if success:
            self.completed += 1
            status_text = f"✓ {message}"
        else:
            self.failed += 1
            status_text = f"✗ {message}"
        
        # Обновляем прогресс
        current = self.completed + self.failed
        self.progress['value'] = current
        self.counter_label.config(text=f"{current} из {self.total}")
        self.status.config(text=status_text)
        
        # Если все файлы обработаны
        if current == self.total:
            self.is_complete = True
            result_text = f"Завершено!\nУспешно: {self.completed}\nОшибок: {self.failed}"
            self.label.config(text=result_text)
            self.close_btn.config(state=tk.NORMAL)
        
        self.top.update()
    
    def close(self):
        """Закрывает окно прогресса"""
        if self.is_complete or self.completed + self.failed == self.total:
            self.top.destroy()
        else:
            # Если печать еще не завершена, спрашиваем подтверждение
            if messagebox.askyesno("Подтверждение", 
                                   "Печать еще не завершена. Прервать?"):
                self.top.destroy()

def start_print_job(print_func, files):
    """Запускает печать в отдельном потоке"""
    if not files:
        messagebox.showwarning("Ошибка", "Нет файлов для печати")
        return
    
    progress_window = PrintProgressWindow(root, len(files))
    
    def worker():
        for file_path in files:
            if not os.path.exists(file_path):
                progress_window.update(False, f"Файл не найден: {os.path.basename(file_path)}")
                continue
            
            # Вызываем функцию печати
            print_func(file_path, progress_window.update)
            
            # Небольшая пауза между заданиями печати
            time.sleep(0.5)
    
    # Запускаем в отдельном потоке
    thread = threading.Thread(target=worker, daemon=True)
    thread.start()


# Список для хранения отобранных строк
selected_rows = []

def format_product_type(product_type):
    """Форматирует тип товара: первая буква заглавная, остальные строчные"""
    if not product_type or not isinstance(product_type, str):
        return product_type
    return product_type.strip().capitalize()

def update_selected_rows_listbox():
    """Обновляет список выбранных организаций с группировкой по компании"""
    selected_rows_listbox.delete(0, tk.END)
    
    # Группируем организации
    orgs = {}
    for name, inn, email, product_type in selected_rows:
        key = (name, inn, email)
        if key not in orgs:
            orgs[key] = set()
        if product_type and product_type != 'не указан':
            types = [format_product_type(t.strip()) for t in str(product_type).split(',')]
            orgs[key].update(types)
    
    # Выводим в список
    for idx, ((name, inn, email), types) in enumerate(orgs.items(), 1):
        types_str = ', '.join(sorted(types)) if types else 'не указан'
        entry = (
            f"{idx}. {name[:50]}{'...' if len(name) > 50 else ''}\n"
            f"   ИНН: {inn if inn else 'не указан'}\n"
            f"   Email: {email if email else 'не указан'}\n"
            f"   Типы товаров: {types_str[:100]}{'...' if len(types_str) > 100 else ''}"
        )
        selected_rows_listbox.insert(tk.END, entry)

def fuzzy_search():
    """Выполняет поиск по наименованию или типу товара (без учета регистра)"""
    query = search_query_var.get().strip().lower()
    search_type = search_type_var.get()
    
    if not query and search_type == "name":
        messagebox.showwarning("Ошибка", "Введите запрос для поиска.")
        return

    output_file = output_file_var_fz.get()
    if not output_file or not os.path.exists(output_file):
        messagebox.showwarning("Ошибка", "Файл output.xlsx не выбран или не существует.")
        return

    try:
        df = pd.read_excel(output_file)
        search_results_listbox.delete(0, tk.END)
        search_results_listbox.config(selectmode=tk.MULTIPLE)
        
        if search_type == "name":
            if 'Наименование' not in df.columns:
                messagebox.showwarning("Ошибка", "В файле отсутствует колонка 'Наименование'.")
                return
                
            # Ищем по наименованию организации (без учета регистра)
            mask = df['Наименование'].str.lower().str.contains(query, na=False, regex=False)
            results = df[mask]['Наименование'].tolist()
            
            # Добавляем в список результатов
            for org_name in results[:50]:
                search_results_listbox.insert(tk.END, org_name)
                
        elif search_type == "product_type":
            if 'Тип товара' not in df.columns:
                messagebox.showwarning("Ошибка", "В файле отсутствует колонка 'Тип товара'.")
                return
            
            # Если запрос пустой, показываем все типы товаров
            if not query:
                show_all_product_types()
                return
            
            # Ищем по типу товара (без учета регистра)
            results = set()
            for _, row in df.iterrows():
                product_types = str(row.get('Тип товара', '')).strip()
                if not product_types:
                    continue
                
                # Разбиваем на отдельные типы товаров
                for pt in product_types.split(','):
                    pt_clean = pt.strip()
                    if pt_clean and query in pt_clean.lower():
                        formatted_pt = format_product_type(pt_clean)
                        org_name = row.get('Наименование', '')
                        results.add(f"{formatted_pt} | Организация: {org_name}")
            
            for result in sorted(results)[:50]:
                search_results_listbox.insert(tk.END, result)
            
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")
        logging.error(f"Ошибка поиска: {str(e)}")

def add_selected_row():
    """Добавляет выбранные организации в список"""
    selected_indices = search_results_listbox.curselection()
    if not selected_indices:
        messagebox.showwarning("Ошибка", "Выберите организации из списка")
        return

    try:
        df = pd.read_excel(output_file_var_fz.get())
        search_type = search_type_var.get()
        added_count = 0

        for index in selected_indices:
            selected_value = search_results_listbox.get(index)
            
            if search_type == "product_type":
                # Проверяем, содержит ли строка разделитель " | Организация: "
                if " | Организация: " in selected_value:
                    # Извлекаем тип товара и организацию из строки
                    product_type, org_name = selected_value.split(" | Организация: ", 1)
                    
                    # Ищем организацию с таким наименованием (без учета регистра)
                    mask = df['Наименование'].str.lower() == org_name.lower()
                    matching_rows = df[mask]
                else:
                    # Если это просто тип товара без организации (из show_all_product_types)
                    product_type = selected_value
                    # Ищем все организации с таким типом товара
                    mask = df['Тип товара'].str.lower().str.contains(product_type.lower(), na=False, regex=False)
                    matching_rows = df[mask]
                
                for _, row in matching_rows.iterrows():
                    org_data = (
                        row['Наименование'],
                        str(row.get('ИНН', 'не указан')),
                        str(row.get('Электронная почта', 'не указан')),
                        str(row.get('Тип товара', 'не указан'))
                    )
                    
                    # Проверяем, есть ли уже такая организация в списке
                    is_duplicate = False
                    for existing in selected_rows:
                        if existing[0] == org_data[0] and existing[1] == org_data[1]:
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        selected_rows.append(org_data)
                        added_count += 1
            
            elif search_type == "name":
                # Для наименования ищем точное совпадение
                mask = df['Наименование'] == selected_value
                matching_rows = df[mask]
                
                for _, row in matching_rows.iterrows():
                    org_data = (
                        row['Наименование'],
                        str(row.get('ИНН', 'не указан')),
                        str(row.get('Электронная почта', 'не указан')),
                        str(row.get('Тип товара', 'не указан'))
                    )
                    
                    # Проверяем на дубликаты
                    is_duplicate = False
                    for existing in selected_rows:
                        if existing[0] == org_data[0] and existing[1] == org_data[1]:
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        selected_rows.append(org_data)
                        added_count += 1

        update_selected_rows_listbox()
        
        if added_count > 0:
            messagebox.showinfo("Успех", f"Добавлено {added_count} организаций")
        else:
            messagebox.showinfo("Информация", "Все выбранные организации уже добавлены")
            
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при добавлении: {str(e)}")
        logging.error(f"Ошибка добавления: {str(e)}")

def remove_selected_row():
    """Удаляет выбранные строки из списка"""
    selected_index = selected_rows_listbox.curselection()
    if not selected_index:
        messagebox.showwarning("Ошибка", "Выберите строку для удаления.")
        return
    
    for index in reversed(selected_index):
        selected_rows.pop(index)
    
    update_selected_rows_listbox()

def search_by_name():
    """Поиск организаций по наименованию"""
    query = search_query_var.get().strip()
    if not query:
        messagebox.showwarning("Ошибка", "Введите название организации")
        search_entry.focus_set()
        return

    search_type_var.set("name")
    fuzzy_search()

def show_all_product_types():
    """Показывает все уникальные типы товаров из файла с форматированием (первая буква заглавная)"""
    output_file = output_file_var_fz.get()
    if not output_file or not os.path.exists(output_file):
        messagebox.showwarning("Ошибка", "Файл output.xlsx не выбран или не существует.")
        return

    try:
        df = pd.read_excel(output_file)
        if 'Тип товара' not in df.columns:
            messagebox.showwarning("Ошибка", "В файле отсутствует колонка 'Тип товара'.")
            return

        # Получаем все уникальные типы товаров с форматированием
        product_types = set()
        for types in df['Тип товара'].dropna():
            if isinstance(types, str):
                for pt in types.split(','):
                    cleaned_pt = pt.strip()
                    if cleaned_pt:
                        product_types.add(format_product_type(cleaned_pt))

        if not product_types:
            messagebox.showinfo("Информация", "В файле не найдено типов товаров.")
            return

        # Очищаем и заполняем список результатов
        search_results_listbox.delete(0, tk.END)
        search_results_listbox.config(selectmode=tk.MULTIPLE)
        for pt in sorted(product_types):
            search_results_listbox.insert(tk.END, pt)

        search_type_var.set("product_type")
        messagebox.showinfo("Информация", f"Найдено {len(product_types)} типов товаров.\nВы можете выбрать несколько типов для добавления.")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")
        logging.error(f"Ошибка показа типов товаров: {str(e)}")

def generate_documents():
    """
    Создает документы Word для всех выбранных организаций,
    объединяя типы товаров для одинаковых компаний и вставляя номера запросов.
    Также создает два текстовых файла:
    - "Адреса без номеров.txt" - названия организаций как в output.xls
    - "Адреса запросов.txt" - названия организаций с номером запроса в формате "№XXX Название"
    """
    try:
        # Проверка наличия данных
        if not selected_rows:
            messagebox.showerror("Ошибка", "Нет выбранных организаций")
            return

        template_path = template_file_var_fz.get()
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("Ошибка", "Шаблон документа не найден")
            return

        save_folder = save_folder_var_fz.get()
        if not save_folder:
            messagebox.showerror("Ошибка", "Не выбрана папка для сохранения")
            return

        # Получаем начальный номер запроса
        try:
            start_number = int(start_number_var.get())
        except ValueError:
            start_number = 1

        # Создаем папку для документов
        output_folder = os.path.join(save_folder, "Итоговые_документы")
        os.makedirs(output_folder, exist_ok=True)

        # Группируем организации по названию, ИНН и email
        orgs_dict = {}
        for name, inn, email, product_type in selected_rows:
            key = (name, inn, email)
            if key not in orgs_dict:
                orgs_dict[key] = set()
            
            # Типы товаров больше не нужны в реквизитах, но сохраняем для группировки
            if product_type and product_type != 'не указан':
                types = [t.strip() for t in str(product_type).split(',')]
                orgs_dict[key].update(types)

        # Определяем максимальное количество цифр в номере
        total_count = len(orgs_dict)
        max_number = start_number + total_count - 1
        number_digits = len(str(max_number))
        
        # Создаем документы для каждой уникальной организации
        success_count = 0
        current_number = start_number
        
        # Списки для текстовых файлов
        addresses_without_numbers = []  # Для файла "Адреса без номеров.txt"
        addresses_with_numbers = []      # Для файла "Адреса запросов.txt" (с номерами)
        
        for (name, inn, email), product_types in orgs_dict.items():
            try:
                # Форматируем номер с лидирующими нулями
                formatted_number = str(current_number).zfill(number_digits)
                
                # Создаем временный документ для реквизитов
                temp_doc = Document()
                
                # Форматируем реквизиты (БЕЗ ТИПОВ ТОВАРОВ)
                p = temp_doc.add_paragraph()
                p.alignment = 1  # Выравнивание по центру
                
                # Наименование организации
                name_run = p.add_run(f"\n\n{name if name else 'не указано'}\n\n")
                name_run.bold = True
                
                # ИНН
                p.add_run("ИНН: ").bold = True
                p.add_run(f"{inn if inn else 'не указан'}\n")
                
                # Email
                p.add_run("E-mail: ").bold = True
                p.add_run(f"{email if email else 'не указан'}\n")
                
                # Типы товаров УДАЛЕНЫ из реквизитов
                p.add_run("\n\n")

                # Открываем основной шаблон
                doc = Document(template_path)
                
                # Вставляем реквизиты в начало документа
                for element in temp_doc.element.body:
                    doc.element.body.insert(0, element)

                # Функция для замены текста во всех элементах документа
                def replace_text_in_paragraph(paragraph):
                    if '№' in paragraph.text and '/07' in paragraph.text:
                        old_text = paragraph.text
                        # Заменяем "№ /07" на "№ {номер}/07"
                        if '№ /07' in old_text:
                            new_text = old_text.replace('№ /07', f'№ {formatted_number}/07')
                        else:
                            # Альтернативные варианты
                            new_text = re.sub(r'№\s*/07', f'№ {formatted_number}/07', old_text)
                            new_text = re.sub(r'№\s*(\d*)/07', f'№ {formatted_number}/07', new_text)
                        
                        if new_text != old_text:
                            paragraph.text = new_text
                            logging.info(f"Заменен номер в параграфе: {old_text[:30]}... -> {new_text[:30]}...")

                # Обрабатываем все параграфы в документе
                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph)

                # Обрабатываем все таблицы в документе
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph)

                # Обрабатываем верхние и нижние колонтитулы (если есть)
                try:
                    for section in doc.sections:
                        # Верхний колонтитул
                        header = section.header
                        if header:
                            for paragraph in header.paragraphs:
                                replace_text_in_paragraph(paragraph)
                        
                        # Нижний колонтитул
                        footer = section.footer
                        if footer:
                            for paragraph in footer.paragraphs:
                                replace_text_in_paragraph(paragraph)
                except Exception as e:
                    logging.warning(f"Не удалось обработать колонтитулы: {str(e)}")

                # Генерируем имя файла с номером запроса в начале (с лидирующими нулями)
                # Очищаем наименование от недопустимых символов
                clean_name = re.sub(r'[\\/*?:"<>|]', '_', name).strip()
                
                # Формируем имя файла: "#Номер_запроса Наименование_организации.docx"
                # Заменяем пробелы на подчеркивания в наименовании для читаемости
                name_for_filename = clean_name.replace(' ', '_')
                file_name = f"#{formatted_number} {name_for_filename}.docx"
                
                # Если имя файла получилось слишком длинным (больше 200 символов), обрезаем
                if len(file_name) > 200:
                    # Оставляем номер и часть названия
                    max_name_length = 200 - len(f"#{formatted_number} .docx")
                    name_for_filename = name_for_filename[:max_name_length]
                    file_name = f"#{formatted_number} {name_for_filename}.docx"
                
                file_path = os.path.join(output_folder, file_name)
                
                # Проверяем, не существует ли уже файл с таким именем
                counter = 1
                original_file_path = file_path
                final_file_name = file_name
                while os.path.exists(file_path):
                    # Если файл существует, добавляем счетчик в конец имени
                    base_name = f"#{formatted_number} {name_for_filename}"
                    final_file_name = f"{base_name}_{counter}.docx"
                    file_path = os.path.join(output_folder, final_file_name)
                    counter += 1
                
                # Сохраняем документ
                doc.save(file_path)
                
                # Добавляем информацию в списки для текстовых файлов
                addresses_without_numbers.append(name)
                # Для файла "Адреса запросов" добавляем номер со знаком № (с лидирующими нулями) и название организации
                addresses_with_numbers.append(f"№{formatted_number} {name}")
                
                success_count += 1
                logging.info(f"Создан документ {final_file_name} с номером {formatted_number}")
                current_number += 1

                root.update_idletasks()

            except Exception as e:
                logging.error(f"Ошибка при создании документа для {name}: {str(e)}")
                continue

        # Создаем текстовые файлы после успешного создания документов
        if success_count > 0:
            try:
                # Файл "Адреса без номеров.txt"
                addresses_without_numbers_path = os.path.join(output_folder, "Адреса без номеров.txt")
                with open(addresses_without_numbers_path, 'w', encoding='utf-8') as f:
                    for address in addresses_without_numbers:
                        f.write(f"{address}\n")
                logging.info(f"Создан файл {addresses_without_numbers_path}")
                
                # Файл "Адреса запросов.txt" - с номерами (№) и названиями организаций (с лидирующими нулями)
                addresses_with_numbers_path = os.path.join(output_folder, "Адреса запросов.txt")
                with open(addresses_with_numbers_path, 'w', encoding='utf-8') as f:
                    for item in addresses_with_numbers:
                        f.write(f"{item}\n")
                logging.info(f"Создан файл {addresses_with_numbers_path}")
                
            except Exception as e:
                logging.error(f"Ошибка при создании текстовых файлов: {str(e)}")
                messagebox.showwarning("Предупреждение", 
                                     f"Документы созданы, но не удалось создать текстовые файлы:\n{str(e)}")

            # Показываем результат
            messagebox.showinfo(
                "Успех",
                f"Успешно создано {success_count} из {total_count} документов.\n"
                f"Номера запросов: с {start_number} по {max_number}\n"
                f"Формат номеров: {number_digits} знаков с лидирующими нулями\n"
                f"Созданы файлы:\n"
                f"- Адреса без номеров.txt\n"
                f"- Адреса запросов.txt\n"
                f"Папка: {output_folder}"
            )
            os.startfile(output_folder)
        else:
            messagebox.showerror(
                "Ошибка",
                "Не удалось создать ни одного документа. Проверьте log.txt"
            )

    except Exception as e:
        messagebox.showerror(
            "Критическая ошибка",
            f"Произошла ошибка:\n{str(e)}\n\nПодробности в log.txt"
        )
        logging.error(f"Критическая ошибка: {str(e)}", exc_info=True)

def clean_text(text):
    """Очистка текста от лишних символов."""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', str(text).strip())

def clean_inn(inn):
    """Очистка и проверка ИНН."""
    if not inn:
        return "не указан"
    
    inn_str = re.sub(r'\D', '', str(inn))
    return inn_str if len(inn_str) in (10, 12) else "не указан"

def clean_email(email):
    """Проверка корректности email."""
    if not email:
        return "не указан"
    
    email_str = str(email).strip()
    return email_str if re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', email_str) else "не указан"

def generate_filename(name, index):
    """Генерация корректного имени файла."""
    clean_name = re.sub(r'[\\/*?:"<>|]', '_', name)[:50]
    return f"{clean_name or f'Документ_{index}'}.docx"

def get_unique_filename(folder, filename):
    """Получение уникального имени файла."""
    base, ext = os.path.splitext(filename)
    counter = 1
    file_path = os.path.join(folder, filename)
    
    while os.path.exists(file_path):
        file_path = os.path.join(folder, f"{base}_{counter}{ext}")
        counter += 1
    
    return file_path

def save_settings():
    """Сохраняет все настройки программы в JSON-файл"""
    settings = {
        'working_folder': working_folder_var.get(),
        'compare_inn': compare_inn_var.get(),
        'compare_inn_value': compare_inn_value_var.get(),
        'output_file': output_file_var.get(),
        'template_file': template_file_var.get(),
        'replace_from': replace_from_var.get(),
        'replace_to': replace_to_var.get(),
        'output_file_fz': output_file_var_fz.get(),
        'template_file_fz': template_file_var_fz.get(),
        'save_folder_fz': save_folder_var_fz.get(),
        'print_folder': print_folder_var.get(),
        'start_number': start_number_var.get()
    }
    
    try:
        with open('settings.json', 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("Успех", "Настройки успешно сохранены!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить настройки:\n{str(e)}")

def load_settings():
    """Загружает все настройки программы из JSON-файла"""
    try:
        if os.path.exists('settings.json'):
            with open('settings.json', 'r', encoding='utf-8') as f:
                settings = json.load(f)
                
            working_folder_var.set(settings.get('working_folder', ''))
            compare_inn_var.set(settings.get('compare_inn', False))
            compare_inn_value_var.set(settings.get('compare_inn_value', ''))
            output_file_var.set(settings.get('output_file', ''))
            template_file_var.set(settings.get('template_file', ''))
            replace_from_var.set(settings.get('replace_from', ''))
            replace_to_var.set(settings.get('replace_to', ''))
            output_file_var_fz.set(settings.get('output_file_fz', ''))
            template_file_var_fz.set(settings.get('template_file_fz', ''))
            save_folder_var_fz.set(settings.get('save_folder_fz', ''))
            print_folder_var.set(settings.get('print_folder', ''))
            start_number_var.set(settings.get('start_number', '1'))
            
            messagebox.showinfo("Успех", "Настройки успешно загружены!")
        else:
            messagebox.showinfo("Информация", "Файл настроек не найден. Будет создан новый при сохранении.")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось загрузить настройки:\n{str(e)}")

def create_requisites_file():
    """Создает текстовый файл с реквизитами отобранных организаций"""
    if not selected_rows:
        messagebox.showwarning("Ошибка", "Нет отобранных организаций")
        return

    save_folder = save_folder_var_fz.get()
    if not save_folder:
        messagebox.showwarning("Ошибка", "Не выбрана папка для сохранения")
        return

    try:
        output_folder = os.path.join(save_folder, "Итоговые_документы")
        os.makedirs(output_folder, exist_ok=True)
        
        file_path = os.path.join(output_folder, "Реквизиты.txt")
        
        # Группируем организации для более компактного вывода
        orgs_dict = {}
        for name, inn, email, product_type in selected_rows:
            key = (name, inn, email)
            if key not in orgs_dict:
                orgs_dict[key] = set()
            if product_type and product_type != 'не указан':
                types = [format_product_type(t.strip()) for t in str(product_type).split(',')]
                orgs_dict[key].update(types)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            for i, ((name, inn, email), types) in enumerate(orgs_dict.items(), 1):
                f.write(f"{i}. {name}\n")
                f.write(f"   ИНН: {inn if inn else 'не указан'}\n")
                f.write(f"   Email: {email if email else 'не указан'}\n")
                
                if types:
                    f.write(f"   Типы товаров: {', '.join(sorted(types))}\n\n")
                else:
                    f.write(f"   Типы товаров: не указаны\n\n")
        
        messagebox.showinfo("Успех", f"Файл с реквизитами успешно создан:\n{file_path}")
        os.startfile(output_folder)
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось создать файл:\n{str(e)}")
        logging.error(f"Ошибка создания файла реквизитов: {str(e)}")

# Функция для центрирования окна
def center_window(window, width, height):
    """Центрирует окно на экране"""
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    window.geometry(f"{width}x{height}+{x}+{y}")

# Создаем графический интерфейс
root = tk.Tk()
root.title("Коммерческие предложения в один клик v1.1")
root.geometry("900x700")

# Центрируем главное окно
center_window(root, 900, 700)

# Проверим, установлен ли Word
if not is_word_installed():
    messagebox.showwarning(
        "Предупреждение", 
        "Microsoft Word не обнаружен!\n"
        "Функции печати будут недоступны.\n"
        "Установите Microsoft Office для использования этой функции."
    )

# Переменные для хранения путей
compare_inn_value_var = tk.StringVar()
compare_inn_var = tk.BooleanVar()
output_file_var = tk.StringVar()
output_file_var_fz = tk.StringVar()
replace_from_var = tk.StringVar()
replace_to_var = tk.StringVar()
save_folder_var_fz = tk.StringVar()
search_query_var = tk.StringVar()
template_file_var = tk.StringVar()
template_file_var_fz = tk.StringVar()
working_folder_var = tk.StringVar()
print_folder_var = tk.StringVar()
start_number_var = tk.StringVar(value="1")
search_type_var = tk.StringVar(value="name")

# Создаем вкладки
notebook = ttk.Notebook(root)
notebook.pack(fill='both', expand=True, padx=10, pady=10)

# Вкладка для создания output.xlsx
tab1 = ttk.Frame(notebook)
notebook.add(tab1, text="Создать output.xlsx")

# Центрируем содержимое вкладки 1
tab1.grid_columnconfigure(0, weight=1)
tab1.grid_columnconfigure(1, weight=1)
tab1.grid_columnconfigure(2, weight=1)

tk.Label(tab1, text="Рабочая папка:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab1, textvariable=working_folder_var, width=50).grid(row=0, column=1, padx=5, pady=5)
tk.Button(tab1, text="Выбрать", command=select_working_folder).grid(row=0, column=2, padx=5, pady=5, sticky='w')

tk.Checkbutton(tab1, text="Искать ИНН через Яндекс, если не найден или совпадает с указанным", variable=compare_inn_var).grid(row=1, column=0, columnspan=3, padx=5, pady=5)

tk.Label(tab1, text="ИНН для сравнения:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab1, textvariable=compare_inn_value_var, width=50).grid(row=2, column=1, padx=5, pady=5)

progress_bar = ttk.Progressbar(tab1, orient="horizontal", length=400, mode="determinate")
progress_bar.grid(row=3, column=0, columnspan=3, padx=5, pady=5)

status_label = tk.Label(tab1, text="Ожидание начала обработки")
status_label.grid(row=4, column=0, columnspan=3, padx=5, pady=5)

tk.Button(tab1, text="Создать output.xlsx", command=lambda: create_output_file(progress_bar, status_label)).grid(row=5, column=1, padx=5, pady=20)

# Вкладка для создания документов
tab2 = ttk.Frame(notebook)
notebook.add(tab2, text="Формирование шаблонов из таблицы")

# Центрируем содержимое вкладки 2
tab2.grid_columnconfigure(0, weight=1)
tab2.grid_columnconfigure(1, weight=1)
tab2.grid_columnconfigure(2, weight=1)

tk.Label(tab2, text="Рабочая папка:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab2, textvariable=working_folder_var, width=50).grid(row=0, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_working_folder).grid(row=0, column=2, padx=5, pady=5, sticky='w')

tk.Label(tab2, text="Файл output.xlsx:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab2, textvariable=output_file_var, width=50).grid(row=1, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_output_file).grid(row=1, column=2, padx=5, pady=5, sticky='w')

tk.Label(tab2, text="Файл шаблона:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab2, textvariable=template_file_var, width=50).grid(row=2, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_template_file).grid(row=2, column=2, padx=5, pady=5, sticky='w')

tk.Label(tab2, text="Текст для замены:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab2, textvariable=replace_from_var, width=50).grid(row=3, column=1, padx=5, pady=5)

tk.Label(tab2, text="Новый текст:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab2, textvariable=replace_to_var, width=50).grid(row=4, column=1, padx=5, pady=5)

progress_bar2 = ttk.Progressbar(tab2, orient="horizontal", length=400, mode="determinate")
progress_bar2.grid(row=5, column=0, columnspan=3, padx=5, pady=5)

status_label2 = tk.Label(tab2, text="Ожидание начала обработки")
status_label2.grid(row=6, column=0, columnspan=3, padx=5, pady=5)

tk.Button(tab2, text="Создать документы", command=lambda: process_files(progress_bar2, status_label2)).grid(row=7, column=1, padx=5, pady=20)

# Вкладка "Формирование запросов"
tab3 = ttk.Frame(notebook)
notebook.add(tab3, text="Формирование запросов")

# Центрируем содержимое вкладки 3
for i in range(4):
    tab3.grid_columnconfigure(i, weight=1)

# Элементы управления
controls_frame = tk.Frame(tab3)
controls_frame.grid(row=0, column=0, columnspan=4, padx=5, pady=5, sticky='n')

# Центрируем содержимое фрейма управления
controls_frame.grid_columnconfigure(1, weight=1)

# Файл шаблона
tk.Label(controls_frame, text="Файл шаблона:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
tk.Entry(controls_frame, textvariable=template_file_var_fz, width=50).grid(row=0, column=1, padx=5, pady=5)
tk.Button(
    controls_frame, 
    text="Выбрать", 
    command=lambda: template_file_var_fz.set(filedialog.askopenfilename(filetypes=[("Word files", "*.docx")]))
).grid(row=0, column=2, padx=5, pady=5, sticky='w')

# Файл output.xlsx
tk.Label(controls_frame, text="Файл output.xlsx:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
tk.Entry(controls_frame, textvariable=output_file_var_fz, width=50).grid(row=1, column=1, padx=5, pady=5)
tk.Button(
    controls_frame, 
    text="Выбрать", 
    command=lambda: output_file_var_fz.set(filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")]))
).grid(row=1, column=2, padx=5, pady=5, sticky='w')

# Папка для сохранения
tk.Label(controls_frame, text="Папка для сохранения:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
tk.Entry(controls_frame, textvariable=save_folder_var_fz, width=50).grid(row=2, column=1, padx=5, pady=5)
tk.Button(
    controls_frame, 
    text="Выбрать", 
    command=lambda: save_folder_var_fz.set(filedialog.askdirectory())
).grid(row=2, column=2, padx=5, pady=5, sticky='w')

# Начальный номер запроса
tk.Label(controls_frame, text="Начальный номер запроса:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
tk.Entry(controls_frame, textvariable=start_number_var, width=50).grid(row=3, column=1, padx=5, pady=5)

# Поисковая строка
search_frame = tk.Frame(tab3)
search_frame.grid(row=4, column=0, columnspan=4, padx=5, pady=5, sticky='n')

# Центрируем содержимое фрейма поиска
search_frame.grid_columnconfigure(1, weight=1)

tk.Label(search_frame, text="Поисковой запрос:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
search_entry = tk.Entry(search_frame, textvariable=search_query_var, width=50)
search_entry.grid(row=0, column=1, padx=5, pady=5)
search_entry.focus_set()

# Кнопки поиска
search_buttons_frame = tk.Frame(tab3)
search_buttons_frame.grid(row=5, column=0, columnspan=4, padx=5, pady=5)

tk.Button(search_buttons_frame, text="Поиск по наименованию", 
          command=search_by_name).pack(side=tk.LEFT, padx=5)
tk.Button(search_buttons_frame, text="Показать все типы товаров", 
          command=show_all_product_types).pack(side=tk.LEFT, padx=5)

# Список результатов поиска с кнопкой очистки
results_frame = tk.Frame(tab3)
results_frame.grid(row=6, column=0, columnspan=4, padx=5, pady=5, sticky='nsew')

# Настраиваем веса для растягивания
tab3.grid_rowconfigure(6, weight=1)
tab3.grid_rowconfigure(8, weight=1)

search_results_listbox = Listbox(
    results_frame, 
    selectmode=tk.MULTIPLE,
    width=70,
    height=8,
    font=('Tahoma', 9)
)
search_results_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

scrollbar = ttk.Scrollbar(results_frame, orient="vertical", command=search_results_listbox.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
search_results_listbox.config(yscrollcommand=scrollbar.set)

clear_results_btn = tk.Button(
    results_frame,
    text="Очистить",
    command=lambda: search_results_listbox.delete(0, tk.END),
    width=8
)
clear_results_btn.pack(side=tk.RIGHT, padx=(5,0))

# Кнопка добавления выбранного
add_button_frame = tk.Frame(tab3)
add_button_frame.grid(row=7, column=0, columnspan=4, padx=5, pady=5)

tk.Button(add_button_frame, text="Добавить выбранное", 
          command=add_selected_row).pack()

# Список отобранных организаций с кнопкой очистки
selected_frame = tk.Frame(tab3)
selected_frame.grid(row=8, column=0, columnspan=4, padx=5, pady=5, sticky='nsew')

selected_rows_listbox = Listbox(
    selected_frame,
    selectmode=tk.MULTIPLE,
    width=70,
    height=8,
    font=('Tahoma', 9),
    bg='#f0f0f0',
    selectbackground='#a6d8ff'
)
selected_rows_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

scrollbar_selected = ttk.Scrollbar(selected_frame, orient="vertical", command=selected_rows_listbox.yview)
scrollbar_selected.pack(side=tk.RIGHT, fill=tk.Y)
selected_rows_listbox.config(yscrollcommand=scrollbar_selected.set)

clear_selected_btn = tk.Button(
    selected_frame,
    text="Очистить",
    command=lambda: [selected_rows.clear(), update_selected_rows_listbox()],
    width=8,
    bg='#ffdddd'
)
clear_selected_btn.pack(side=tk.RIGHT, padx=(5,0))

# Кнопка удаления выбранного
remove_button_frame = tk.Frame(tab3)
remove_button_frame.grid(row=9, column=0, columnspan=4, padx=5, pady=5)

tk.Button(remove_button_frame, text="Удалить выбранное", 
          command=remove_selected_row).pack()

# Кнопки формирования документов
doc_buttons_frame = tk.Frame(tab3)
doc_buttons_frame.grid(row=10, column=0, columnspan=4, padx=5, pady=10)

tk.Button(doc_buttons_frame, 
          text="Сформировать текстовый файл с реквизитами", 
          command=create_requisites_file,
          bg="#e6e6fa").pack(side=tk.LEFT, padx=5)

tk.Button(doc_buttons_frame, 
          text="Сформировать документы", 
          command=generate_documents).pack(side=tk.LEFT, padx=5)

# Привязка горячих клавиш
root.bind('<Return>', lambda e: search_by_name())

# Вкладка "Печать"
tab4 = ttk.Frame(notebook)
notebook.add(tab4, text="Печать")

# Центрируем содержимое вкладки 4
tab4.grid_columnconfigure(0, weight=1)
tab4.grid_columnconfigure(1, weight=1)
tab4.grid_columnconfigure(2, weight=1)

tk.Label(tab4, text="Папка с документами:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
tk.Entry(tab4, textvariable=print_folder_var, width=50).grid(row=0, column=1, padx=5, pady=5)
tk.Button(tab4, text="Выбрать", command=lambda: print_folder_var.set(filedialog.askdirectory())).grid(row=0, column=2, padx=5, pady=5, sticky='w')

# Добавляем кнопки печати с подтверждением
tk.Button(tab4, text="Печать первых страниц", 
          command=print_first_pages,
          bg="#e6f3ff").grid(row=1, column=1, padx=5, pady=10, sticky='ew')

tk.Button(tab4, text="Печать всех документов", 
          command=print_all_documents,
          bg="#e6f3ff").grid(row=2, column=1, padx=5, pady=10, sticky='ew')

# Информационная метка
info_label = tk.Label(tab4, text="Для печати используется принтер по умолчанию", 
                     font=('Tahoma', 8), fg='gray')
info_label.grid(row=3, column=1, padx=5, pady=5)

# Кнопки сохранения/загрузки настроек внизу окна
settings_frame = tk.Frame(root)
settings_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=5)

# Центрируем кнопки настроек
settings_frame.grid_columnconfigure(0, weight=1)
settings_frame.grid_columnconfigure(1, weight=1)

tk.Button(settings_frame, text="Сохранить настройки", command=save_settings).grid(row=0, column=0, padx=5, pady=5, sticky='e')
tk.Button(settings_frame, text="Загрузить настройки", command=load_settings).grid(row=0, column=1, padx=5, pady=5, sticky='w')

# Запуск основного цикла GUI
if __name__ == "__main__":
    load_settings()
    root.mainloop()