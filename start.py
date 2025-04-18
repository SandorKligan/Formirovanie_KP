from bs4 import BeautifulSoup
from docx import Document
from fuzzywuzzy import process
from tkinter import ttk, filedialog, messagebox, StringVar, Listbox, MULTIPLE, END, Tk, Entry, Button
import glob
import json
import logging
import os
import pandas as pd
import random
import re
import requests
import subprocess
import time
import tkinter as tk
import win32com.client as win32


# Отключаем логирование для win32com
logging.getLogger('win32com').setLevel(logging.WARNING)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    # Лог сохраняется в текущей папке
    filename=os.path.join(os.getcwd(), 'log.txt'),
    filemode='w'
)


def extract_text_from_docx(file_path):
    """Извлекает текст из файла .docx."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return full_text


def extract_text_from_doc(file_path):
    """Извлекает текст из файла .doc."""
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(file_path)
    full_text = []
    for para in doc.Paragraphs:
        full_text.append(para.Range.Text.strip())
    doc.Close()
    word.Quit()
    return full_text


def extract_info(text_lines):
    info = {'Наименование': '', 'ИНН': '', 'Адрес': '',
            'Электронная почта': '', 'Телефон': ''}

    # Убираем пустые строки
    text_lines = [line for line in text_lines if line.strip()]

    # Останавливаем обработку при обнаружении слова "ЗАПРОС" (в любом регистре)
    stop_index = None
    for i, line in enumerate(text_lines):
        if "ЗАПРОС" in line.upper():
            stop_index = i
            break

    # Если слово "ЗАПРОС" не найдено, ищем три пустые строки подряд
    if stop_index is None:
        empty_lines_count = 0
        for i, line in enumerate(text_lines):
            if not line.strip():
                empty_lines_count += 1
                if empty_lines_count == 3:
                    stop_index = i - 2  # Останавливаемся перед тремя пустыми строками
                    break
            else:
                empty_lines_count = 0

    # Если найдено условие остановки, обрезаем текст
    if stop_index is not None:
        text_lines = text_lines[:stop_index]

    # Объединяем строки для упрощения поиска
    combined_text = ' '.join(text_lines)

    # Поиск наименования
    # Если есть ИНН, берем текст до ИНН
    inn_index = next((i for i, line in enumerate(
        text_lines) if 'ИНН' in line), None)
    if inn_index is not None:
        info['Наименование'] = ' '.join(line.strip()
                                        for line in text_lines[:inn_index])
    else:
        # Если ИНН нет, берем первую строку как наименование
        info['Наименование'] = text_lines[0].strip()

    # Удаляем обращения ("Руководителю", "ИП", "Индивидуальный предприниматель", "Индивидуальному предпринимателю", "Директору", "Генеральному директору")
    info['Наименование'] = re.sub(
        r'^(Руководителю|ИП|Индивидуальный предприниматель|Индивидуальному предпринимателю|Директору|Генеральному директору)\s*',
        '', info['Наименование'], flags=re.IGNORECASE
    ).strip()

    # Если наименование начинается с кавычек, добавляем "ООО"
    if info['Наименование'].startswith(('«', '"', "'")):
        info['Наименование'] = f"ООО {info['Наименование']}"

    # Заменяем "Общество с ограниченной ответственностью" на "ООО"
    info['Наименование'] = info['Наименование'].replace(
        "Общество с ограниченной ответственностью", "ООО")

    # Поиск ИНН
    inn_match = re.search(r'ИНН\s*(\d{10,12})', combined_text)
    if inn_match:
        info['ИНН'] = inn_match.group(1).strip()

    # Поиск адреса
    address_match = re.search(
        r'(?:Адрес|Юридический адрес):\s*([\s\S]+?)(?=\sE-mail|\sТелефон|$)', combined_text)
    if address_match:
        info['Адрес'] = address_match.group(1).strip()

    # Поиск всех электронных почт
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', combined_text)
    if emails:
        info['Электронная почта'] = ', '.join(emails)

    # Поиск всех номеров телефонов
    phones = re.findall(r'\+?\d[\d\s\-()]{6,}\d', combined_text)
    if phones:
        info['Телефон'] = ', '.join(phones)

    # Сохраняем всю информацию до последней строки с email
    info['Исходная информация'] = '\n'.join(text_lines)

    return info


def is_valid_inn(inn: str) -> bool:
    if len(inn) == 10:
        weights = [2, 4, 10, 3, 5, 9, 4, 6, 8]
    elif len(inn) == 12:
        weights = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8, 0]
    else:
        return False

    checksum = sum(int(c) * w for c, w in zip(inn[:-1], weights)) % 11 % 10
    return checksum == int(inn[-1])


def get_random_user_agent():
    agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) ...",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) ...",
    ]
    return random.choice(agents)


def get_inn_by_name(organization_name: str, timeout: int = 10) -> str | None:
    """
    Поиск ИНН организации через Яндекс с улучшенной точностью и защитой от блокировки.

    :param organization_name: Название организации.
    :param timeout: Таймаут запроса в секундах.
    :return: ИНН или None, если не найден.
    """
    try:
        # Формирование запроса
        search_query = f"ИНН {organization_name}"
        url = f"https://yandex.ru/search/?text={requests.utils.quote(search_query)}"

        # Случайные заголовки и задержка
        headers = {
            "User-Agent": get_random_user_agent(),
            "Accept-Language": "ru-RU,ru;q=0.9",
        }
        time.sleep(random.uniform(1, 3))

        # Запрос
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()

        # Парсинг
        soup = BeautifulSoup(response.text, 'html.parser')
        texts = soup.find_all(text=re.compile(
            r'ИНН.*?\d{10,12}', re.IGNORECASE))

        for text in texts:
            inn_match = re.search(r'(?<![\d-])(\d{10,12})(?![\d-])', text)
            if inn_match and is_valid_inn(inn_match.group(1)):
                return inn_match.group(1)

        logging.warning(f"ИНН для '{organization_name}' не найден.")
        return None

    except Exception as e:
        logging.error(f"Ошибка: {e}")
        return None


def create_output_file(progress_bar, status_label):
    """Создает файл output.xlsx на основе данных из документов."""
    current_dir = working_folder_var.get()
    if not current_dir:
        messagebox.showerror("Ошибка", "Выберите рабочую папку.")
        return

    doc_files = glob.glob(os.path.join(current_dir, '*.doc')) + \
        glob.glob(os.path.join(current_dir, '*.docx'))
    if not doc_files:
        messagebox.showerror("Ошибка", "В папке нет файлов .doc или .docx.")
        return

    data = []
    total_files = len(doc_files)
    for idx, file_path in enumerate(doc_files):
        try:
            if file_path.endswith('.docx'):
                text_lines = extract_text_from_docx(file_path)
            else:
                text_lines = extract_text_from_doc(file_path)

            info = extract_info(text_lines)

            # Удаляем строку, если наименование начинается с "Запрос", "Добрый", "ЕИС", "Единая" (независимо от регистра)
            if info['Наименование'].lower().startswith(('запрос', 'добрый', 'еис', 'единая')):
                continue

            info['Номер п/п'] = idx + 1

            # Если ИНН не найден или совпадает с указанным значением, ищем через Яндекс
            if compare_inn_var.get() and (not info['ИНН'] or info['ИНН'] == compare_inn_value_var.get()):
                inn = get_inn_by_name(info['Наименование'])
                if inn:
                    info['ИНН'] = inn

            data.append(info)
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {file_path}: {e}")

        # Обновляем прогресс
        progress = (idx + 1) / total_files * 100
        progress_bar['value'] = progress
        status_label.config(
            text=f"Обработано {idx + 1} из {total_files} файлов")
        root.update_idletasks()

    # Создаем DataFrame и сохраняем в Excel
    df = pd.DataFrame(data)
    columns_order = ['Номер п/п', 'Наименование', 'ИНН', 'Адрес',
                     'Электронная почта', 'Телефон', 'Исходная информация']
    df = df[columns_order]  # Упорядочиваем колонки
    # Сохраняем в текущей папке
    output_path = os.path.join(os.getcwd(), 'output.xlsx')
    df.to_excel(output_path, index=False)
    messagebox.showinfo(
        "Успех", f"Файл output.xlsx успешно создан: {output_path}")


def process_files(progress_bar, status_label):
    """Обрабатывает файлы и создает документы на основе шаблона."""
    output_file = output_file_var.get()
    template_file = template_file_var.get()
    working_folder = working_folder_var.get()
    replace_from = replace_from_var.get()
    replace_to = replace_to_var.get()

    if not output_file or not template_file or not working_folder:
        messagebox.showerror(
            "Ошибка", "Выберите рабочую папку, файл output.xlsx и шаблон.")
        return

    try:
        # Читаем файл output.xlsx (только для чтения)
        df = pd.read_excel(output_file)
        required_columns = ['Наименование', 'ИНН', 'Электронная почта']
        if not all(column in df.columns for column in required_columns):
            messagebox.showerror(
                "Ошибка", f"В файле {output_file} отсутствуют необходимые колонки: {required_columns}")
            return

        # Создаем папку для сохранения документов
        output_folder = os.path.join(working_folder, "Итоговые_документы")
        os.makedirs(output_folder, exist_ok=True)

        # Обрабатываем каждую строку
        total_rows = len(df)
        for index, row in df.iterrows():
            name = row['Наименование']
            inn = row['ИНН']
            email = row['Электронная почта']

            # Создаем новый документ на основе шаблона
            doc = Document(template_file)
            for para in doc.paragraphs:
                if replace_from and replace_to:
                    para.text = para.text.replace(replace_from, replace_to)
                if name in para.text:
                    para.text = para.text.replace(name, name)
                if inn and f"ИНН: {inn}" in para.text:
                    para.text = para.text.replace(f"ИНН: {inn}", f"ИНН: {inn}")
                if email and f"E-mail: {email}" in para.text:
                    para.text = para.text.replace(
                        f"E-mail: {email}", f"E-mail: {email}")

            # Сохраняем документ
            # Ограничиваем длину имени файла
            file_name = re.sub(r'[\\/:*?"<>|]', '_', name)[:50]
            file_path = os.path.join(output_folder, f"{file_name}.docx")
            doc.save(file_path)

            # Обновляем прогресс
            progress = (index + 1) / total_rows * 100
            progress_bar['value'] = progress
            status_label.config(
                text=f"Обработано {index + 1} из {total_rows} строк")
            root.update_idletasks()

        messagebox.showinfo("Успех", "Документы успешно созданы!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")


def select_working_folder():
    """Выбор рабочей папки."""
    folder = filedialog.askdirectory()
    if folder:
        working_folder_var.set(folder)


def select_output_file():
    """Выбор файла output.xlsx."""
    file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file:
        output_file_var.set(file)


def select_template_file():
    """Выбор файла шаблона."""
    file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file:
        template_file_var.set(file)


def is_word_installed():
    try:
        word = win32com.client.Dispatch('Word.Application')
        word.Quit()
        return True
    except Exception as e:
        logging.error(f"Ошибка при проверке Word: {e}")
        return False

# Функция для печати документа через командную строку Word


def print_document(file_path):
    """Печатает документ через командную строку Word."""
    try:
        # Путь к Word (замените на ваш путь)
        word_path = r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE"

        # Команда для печати документа
        command = [word_path, '/q', '/n',
                   '/mFilePrintDefault', '/mFileExit', file_path]

        # Запуск Word
        subprocess.run(command, check=True)
        return True
    except Exception as e:
        logging.error(f"Ошибка при печати файла {file_path}: {e}")
        return False

# Функция для печати первых страниц


def print_first_page(file_path):
    """Печатает первую страницу документа через win32com.client."""
    try:
        # Открываем Word
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        # Открываем документ
        doc = word.Documents.Open(file_path)

        # Печатаем первую страницу
        doc.PrintOut(Range=win32.constants.wdPrintRangeOfPages, Pages="1")

        # Закрываем документ
        doc.Close()

        # Закрываем Word
        word.Quit()
        return True
    except Exception as e:
        logging.error(f"Ошибка при печати файла {file_path}: {e}")
        return False


def print_first_pages():
    folder = print_folder_var.get()
    if not folder:
        messagebox.showwarning("Ошибка", "Выберите папку с документами.")
        return

    try:
        # Получаем список всех файлов .doc и .docx в папке
        doc_files = glob.glob(os.path.join(folder, '*.doc')) + \
            glob.glob(os.path.join(folder, '*.docx'))

        for file_path in doc_files:
            if not print_first_page(file_path):
                logging.error(f"Не удалось распечатать файл: {file_path}")

        messagebox.showinfo("Успех", "Печать первых страниц завершена!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

# Функция для печати всех документов


def print_all_documents():
    folder = print_folder_var.get()
    if not folder:
        messagebox.showwarning("Ошибка", "Выберите папку с документами.")
        return

    try:
        # Получаем список всех файлов .doc и .docx в папке
        doc_files = glob.glob(os.path.join(folder, '*.doc')) + \
            glob.glob(os.path.join(folder, '*.docx'))

        for file_path in doc_files:
            if not print_document(file_path):
                logging.error(f"Не удалось распечатать файл: {file_path}")

        messagebox.showinfo("Успех", "Печать всех документов завершена!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")


# Список для хранения отобранных строк
selected_rows = []

# Функция для нечеткого поиска


def fuzzy_search():
    query = search_query_var.get()
    if not query:
        messagebox.showwarning("Ошибка", "Введите запрос для поиска.")
        return

    output_file = output_file_var_fz.get()
    if not output_file or not os.path.exists(output_file):
        messagebox.showwarning(
            "Ошибка", "Файл output.xlsx не выбран или не существует.")
        return

    try:
        df = pd.read_excel(output_file)
        if 'Наименование' not in df.columns or 'ИНН' not in df.columns or 'Электронная почта' not in df.columns:
            messagebox.showwarning(
                "Ошибка", "В файле output.xlsx отсутствуют необходимые колонки.")
            return

        df['Поисковый текст'] = df['Наименование'] + " | " + \
            df['ИНН'].astype(str) + " | " + df['Электронная почта']
        choices = df['Поисковый текст'].tolist()
        results = process.extract(query, choices, limit=10)

        search_results_listbox.delete(0, tk.END)
        for result in results:
            search_results_listbox.insert(tk.END, result[0])
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

# Функция для добавления строки в список отобранных


def add_selected_row():
    selected_indices = search_results_listbox.curselection()
    if not selected_indices:
        messagebox.showwarning("Ошибка", "Выберите строку из списка.")
        return

    for index in selected_indices:
        selected_text = search_results_listbox.get(index)
        parts = selected_text.split(" | ")

        # Проверяем, что разбиение дало 3 части
        if len(parts) != 3:
            messagebox.showerror(
                "Ошибка", f"Некорректный формат строки: {selected_text}")
            continue

        name = parts[0].strip() if parts[0].strip() else "не указано"
        inn = parts[1].strip() if parts[1].strip() else "не указан"
        email = parts[2].strip() if parts[2].strip() else "не указан"

        selected_rows.append((name, inn, email))

    update_selected_rows_listbox()

# Функция для удаления строки из списка отобранных


def remove_selected_row():
    selected_index = selected_rows_listbox.curselection()
    if not selected_index:
        messagebox.showwarning("Ошибка", "Выберите строку для удаления.")
        return

    selected_rows.pop(selected_index[0])
    update_selected_rows_listbox()

# Функция для обновления списка отобранных строк


def update_selected_rows_listbox():
    selected_rows_listbox.delete(0, tk.END)
    for name, inn, email in selected_rows:
        selected_rows_listbox.insert(tk.END, f"{name} | {inn} | {email}")

# Функция для формирования документов


def generate_documents():
    """
    Создает документы Word с реквизитами в начале, сохраняя исходное форматирование шаблона.
    """
    try:
        # Проверка наличия необходимых данных
        if not selected_rows:
            messagebox.showerror(
                "Ошибка", "Нет выбранных строк для обработки.")
            return

        template_path = template_file_var_fz.get()
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror(
                "Ошибка", "Шаблон документа не выбран или не существует.")
            return

        save_folder = save_folder_var_fz.get()
        if not save_folder:
            messagebox.showerror("Ошибка", "Не выбрана папка для сохранения.")
            return

        # Создаем папку для сохранения
        output_folder = os.path.join(save_folder, "Итоговые_документы")
        os.makedirs(output_folder, exist_ok=True)

        success_count = 0

        for idx, (name, inn, email) in enumerate(selected_rows, 1):
            try:
                # Создаем временный документ для реквизитов
                temp_doc = Document()

                # Добавляем реквизиты во временный документ
                p = temp_doc.add_paragraph()
                p.alignment = 1  # Выравнивание по центру
                p.add_run("\n\n")  # Отступ сверху

                name_run = p.add_run(name if name else "не указано")
                name_run.bold = True
                p.add_run("\n\nИНН: ").bold = False
                p.add_run(inn if inn else "не указан")
                p.add_run("\nE-mail: ")
                p.add_run(email if email else "не указан")
                p.add_run("\n\n")  # Отступ снизу

                # Открываем основной шаблон
                doc = Document(template_path)

                # Вставляем реквизиты в начало основного документа
                for element in temp_doc.element.body:
                    doc.element.body.insert(0, element)

                # Сохранение документа
                file_name = generate_filename(name, idx)
                file_path = get_unique_filename(output_folder, file_name)

                doc.save(file_path)
                success_count += 1

            except Exception as row_error:
                logging.error(
                    f"Ошибка при обработке строки {idx}: {row_error}", exc_info=True)
                continue

        # Итоговое сообщение
        show_result_message(success_count, len(selected_rows), output_folder)

    except Exception as e:
        logging.error(f"Критическая ошибка: {e}", exc_info=True)
        messagebox.showerror(
            "Ошибка",
            f"Произошла критическая ошибка:\n{str(e)}\n\n"
            "Подробности в файле log.txt"
        )


# Вспомогательные функции
def add_centered_paragraph(doc, text, bold=False):
    """Добавляет параграф с центрированным текстом."""
    p = doc.add_paragraph()
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    return p


def clean_text(text):
    """Очистка текста от лишних символов."""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', str(text).strip())


def clean_inn(inn):
    """Очистка и проверка ИНН."""
    if not inn:
        return "не указан"

    inn_str = re.sub(r'\D', '', str(inn))
    return inn_str if len(inn_str) in (10, 12) else "не указан"


def clean_email(email):
    """Проверка корректности email."""
    if not email:
        return "не указан"

    email_str = str(email).strip()
    return email_str if re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', email_str) else "не указан"


def generate_filename(name, index):
    """Генерация корректного имени файла."""
    clean_name = re.sub(r'[\\/*?:"<>|]', '_', name)[:50]
    return f"{clean_name or f'Документ_{index}'}.docx"


def get_unique_filename(folder, filename):
    """Получение уникального имени файла."""
    base, ext = os.path.splitext(filename)
    counter = 1
    file_path = os.path.join(folder, filename)

    while os.path.exists(file_path):
        file_path = os.path.join(folder, f"{base}_{counter}{ext}")
        counter += 1

    return file_path


def show_result_message(success, total, folder):
    """Показ итогового сообщения."""
    if success > 0:
        messagebox.showinfo(
            "Успех",
            f"Успешно создано {success} из {total} документов.\n"
            f"Папка: {folder}"
        )
    else:
        messagebox.showerror(
            "Ошибка",
            "Не удалось создать ни одного документа. Проверьте данные и логи."
        )


def save_settings():
    """Сохраняет все настройки программы в JSON-файл"""
    settings = {
        # Общие настройки
        'working_folder': working_folder_var.get(),

        # Настройки для вкладки "Создать output.xlsx"
        'compare_inn': compare_inn_var.get(),
        'compare_inn_value': compare_inn_value_var.get(),

        # Настройки для вкладки "Формирование шаблонов"
        'output_file': output_file_var.get(),
        'template_file': template_file_var.get(),
        'replace_from': replace_from_var.get(),
        'replace_to': replace_to_var.get(),

        # Настройки для вкладки "Формирование запросов"
        'output_file_fz': output_file_var_fz.get(),
        'template_file_fz': template_file_var_fz.get(),
        'save_folder_fz': save_folder_var_fz.get(),

        # Настройки для вкладки "Печать"
        'print_folder': print_folder_var.get()
    }

    try:
        with open('settings.json', 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("Успех", "Настройки успешно сохранены!")
    except Exception as e:
        messagebox.showerror(
            "Ошибка", f"Не удалось сохранить настройки:\n{str(e)}")


def load_settings():
    """Загружает все настройки программы из JSON-файла"""
    try:
        if os.path.exists('settings.json'):
            with open('settings.json', 'r', encoding='utf-8') as f:
                settings = json.load(f)

            # Общие настройки
            working_folder_var.set(settings.get('working_folder', ''))

            # Настройки для вкладки "Создать output.xlsx"
            compare_inn_var.set(settings.get('compare_inn', False))
            compare_inn_value_var.set(settings.get('compare_inn_value', ''))

            # Настройки для вкладки "Формирование шаблонов"
            output_file_var.set(settings.get('output_file', ''))
            template_file_var.set(settings.get('template_file', ''))
            replace_from_var.set(settings.get('replace_from', ''))
            replace_to_var.set(settings.get('replace_to', ''))

            # Настройки для вкладки "Формирование запросов"
            output_file_var_fz.set(settings.get('output_file_fz', ''))
            template_file_var_fz.set(settings.get('template_file_fz', ''))
            save_folder_var_fz.set(settings.get('save_folder_fz', ''))

            # Настройки для вкладки "Печать"
            print_folder_var.set(settings.get('print_folder', ''))

            messagebox.showinfo("Успех", "Настройки успешно загружены!")
        else:
            messagebox.showinfo(
                "Информация", "Файл настроек не найден. Будет создан новый при сохранении.")
    except Exception as e:
        messagebox.showerror(
            "Ошибка", f"Не удалось загрузить настройки:\n{str(e)}")


# Создаем графический интерфейс
root = tk.Tk()
root.title("Коммерческие предложения в один клик v1.0")

# Переменные для хранения путей
compare_inn_value_var = tk.StringVar()
compare_inn_var = tk.BooleanVar()
output_file_var = tk.StringVar()
output_file_var_fz = tk.StringVar()
print_folder_var = tk.StringVar()
replace_from_var = tk.StringVar()
replace_to_var = tk.StringVar()
save_folder_var_fz = tk.StringVar()
search_query_var = tk.StringVar()
selected_rows = []
template_file_var = tk.StringVar()
template_file_var_fz = tk.StringVar()
working_folder_var = tk.StringVar()


# Создаем вкладки
notebook = ttk.Notebook(root)
notebook.pack(fill='both', expand=True)

# Вкладка для создания output.xlsx
tab1 = ttk.Frame(notebook)
notebook.add(tab1, text="Создать таблицу с реквизитами из имеющихся КП(output.xlsx)")

tk.Label(tab1, text="Папка с КП для анализа:").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(tab1, textvariable=working_folder_var, width=50).grid(
    row=0, column=1, padx=5, pady=5)
tk.Button(tab1, text="Выбрать", command=select_working_folder).grid(
    row=0, column=2, padx=5, pady=5)

tk.Checkbutton(tab1, text="Искать ИНН через Яндекс, если не найден или совпадает с указанным",
               variable=compare_inn_var).grid(row=1, column=0, columnspan=3, padx=5, pady=5)

tk.Label(tab1, text="ИНН для сравнения:").grid(row=2, column=0, padx=5, pady=5)
tk.Entry(tab1, textvariable=compare_inn_value_var,
         width=50).grid(row=2, column=1, padx=5, pady=5)

progress_bar = ttk.Progressbar(
    tab1, orient="horizontal", length=400, mode="determinate")
progress_bar.grid(row=3, column=0, columnspan=3, padx=5, pady=5)

status_label = tk.Label(tab1, text="Ожидание начала обработки")
status_label.grid(row=4, column=0, columnspan=3, padx=5, pady=5)

tk.Button(tab1, text="Создать output.xlsx", command=lambda: create_output_file(
    progress_bar, status_label)).grid(row=5, column=1, padx=5, pady=20)

tk.Button(tab1, text="Сохранить настройки", command=save_settings).grid(
    row=8, column=1, padx=5, pady=5)
tk.Button(tab1, text="Загрузить настройки", command=load_settings).grid(
    row=9, column=1, padx=5, pady=5)
description1=tk.Label(tab1, text=r"Данная вкладка предназначена для формирования таблицы Output(xlsx), содержащей контакты поставщиков. Таблица формируется на основе имеющихся запросов коммерческих предложений.")
description1.grid(row=10, column=0, columnspan=3, padx=10, pady=10)

# Вкладка для создания документов
tab2 = ttk.Frame(notebook)
notebook.add(tab2, text="Формирование шаблонов из таблицы")

tk.Label(tab2, text="Папка с КП для формирования:").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(tab2, textvariable=working_folder_var, width=50).grid(
    row=0, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_working_folder).grid(
    row=0, column=2, padx=5, pady=5)

tk.Label(tab2, text="Файл output.xlsx:").grid(row=1, column=0, padx=5, pady=5)
tk.Entry(tab2, textvariable=output_file_var, width=50).grid(
    row=1, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_output_file).grid(
    row=1, column=2, padx=5, pady=5)

tk.Label(tab2, text="Файл шаблона:").grid(row=2, column=0, padx=5, pady=5)
tk.Entry(tab2, textvariable=template_file_var, width=50).grid(
    row=2, column=1, padx=5, pady=5)
tk.Button(tab2, text="Выбрать", command=select_template_file).grid(
    row=2, column=2, padx=5, pady=5)

tk.Label(tab2, text="Текст для замены:").grid(row=3, column=0, padx=5, pady=5)
tk.Entry(tab2, textvariable=replace_from_var, width=50).grid(
    row=3, column=1, padx=5, pady=5)

tk.Label(tab2, text="Новый текст:").grid(row=4, column=0, padx=5, pady=5)
tk.Entry(tab2, textvariable=replace_to_var, width=50).grid(
    row=4, column=1, padx=5, pady=5)

progress_bar2 = ttk.Progressbar(
    tab2, orient="horizontal", length=400, mode="determinate")
progress_bar2.grid(row=5, column=0, columnspan=3, padx=5, pady=5)

status_label2 = tk.Label(tab2, text="Ожидание начала обработки")
status_label2.grid(row=6, column=0, columnspan=3, padx=5, pady=5)

tk.Button(tab2, text="Создать документы", command=lambda: process_files(
    progress_bar2, status_label2)).grid(row=7, column=1, padx=5, pady=20)
# Создаем вкладку "Формирование запросов"
tab3 = ttk.Frame(notebook)
notebook.add(tab3, text="Формирование запросов")

# Элементы интерфейса для вкладки "Формирование запросов"
tk.Label(tab3, text="Файл шаблона:").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(tab3, textvariable=template_file_var_fz,
         width=50).grid(row=0, column=1, padx=5, pady=5)
tk.Button(tab3, text="Выбрать", command=lambda: template_file_var_fz.set(filedialog.askopenfilename(
    filetypes=[("Word files", "*.docx")]))).grid(row=0, column=2, padx=5, pady=5)

tk.Label(tab3, text="Файл output.xlsx:").grid(row=1, column=0, padx=5, pady=5)
tk.Entry(tab3, textvariable=output_file_var_fz, width=50).grid(
    row=1, column=1, padx=5, pady=5)
tk.Button(tab3, text="Выбрать", command=lambda: output_file_var_fz.set(filedialog.askopenfilename(
    filetypes=[("Excel files", "*.xlsx")]))).grid(row=1, column=2, padx=5, pady=5)

tk.Label(tab3, text="Папка для сохранения:").grid(
    row=2, column=0, padx=5, pady=5)
tk.Entry(tab3, textvariable=save_folder_var_fz, width=50).grid(
    row=2, column=1, padx=5, pady=5)
tk.Button(tab3, text="Выбрать", command=lambda: save_folder_var_fz.set(
    filedialog.askdirectory())).grid(row=2, column=2, padx=5, pady=5)

tk.Label(tab3, text="Поиск по наименованию:").grid(
    row=3, column=0, padx=5, pady=5)
tk.Entry(tab3, textvariable=search_query_var, width=50).grid(
    row=3, column=1, padx=5, pady=5)
tk.Button(tab3, text="Найти", command=fuzzy_search).grid(
    row=3, column=2, padx=5, pady=5)

# Список результатов поиска
search_results_listbox = Listbox(
    tab3, selectmode=tk.MULTIPLE, width=80, height=10)
search_results_listbox.grid(row=4, column=0, columnspan=3, padx=5, pady=5)

# Кнопка для добавления строки в список отобранных
tk.Button(tab3, text="Добавить", command=add_selected_row).grid(
    row=5, column=1, padx=5, pady=5)

# Список отобранных строк
selected_rows_listbox = Listbox(
    tab3, selectmode=tk.MULTIPLE, width=80, height=10)
selected_rows_listbox.grid(row=6, column=0, columnspan=3, padx=5, pady=5)

# Кнопка для удаления строки из списка отобранных
tk.Button(tab3, text="Удалить", command=remove_selected_row).grid(
    row=7, column=1, padx=5, pady=5)

# Кнопка для формирования документов
tk.Button(tab3, text="Сформировать документы", command=generate_documents).grid(
    row=8, column=1, padx=5, pady=20)


# Создаем вкладку "Печать"
tab4 = ttk.Frame(notebook)
notebook.add(tab4, text="Печать")

# Элементы интерфейса для вкладки "Печать"
tk.Label(tab4, text="Папка с документами:").grid(
    row=0, column=0, padx=5, pady=5)
tk.Entry(tab4, textvariable=print_folder_var, width=50).grid(
    row=0, column=1, padx=5, pady=5)
tk.Button(tab4, text="Выбрать", command=lambda: print_folder_var.set(
    filedialog.askdirectory())).grid(row=0, column=2, padx=5, pady=5)

# Кнопка для печати первых страниц
tk.Button(tab4, text="Печать первых страниц", command=print_first_pages).grid(
    row=1, column=1, padx=5, pady=10)

# Кнопка для печати всех документов
tk.Button(tab4, text="Печать всех документов",
          command=print_all_documents).grid(row=2, column=1, padx=5, pady=10)


# Запуск основного цикла GUI
if __name__ == "__main__":
    load_settings()  # Загружаем настройки при старте
    root.mainloop()
